from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_brief_S4():
    document = Document()

    # --- Title ---
    title = document.add_heading('Brief Pédagogique Semaine 4 : Logique & Viz', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Résumé ---
    document.add_heading('1. Résumé de la Semaine', level=1)
    p = document.add_paragraph()
    p.add_run('Cette semaine marque le passage de la simple saisie à l\'analyse. L\'étudiant apprend à faire prendre des décisions à Excel (SI) et à synthétiser l\'information visuellement (Graphiques).')

    # --- Compétences Visées ---
    document.add_heading('2. Compétences Visées', level=1)
    skills = [
        "Comprendre et écrire une fonction logique simple (=SI).",
        "Identifier le type de graphique adapté à la donnée (Comparaison vs Proportion).",
        "Créer, titrer et légender un graphique sur Excel/Sheets.",
        "Lire un graphique (Axe des abscisses et ordonnées)."
    ]
    for skill in skills:
        document.add_paragraph(skill, style='List Bullet')

    # --- Projet Pratique ---
    document.add_heading('3. Description du Projet Pratique', level=1)
    document.add_paragraph('Scénario : "Le Bilan Trimestriel"')
    document.add_paragraph('L\'étudiant reçoit le fichier "dataset_S4_logique_viz.xlsx".')

    document.add_heading('Mission :', level=2)
    missions = [
        "Feuille Notes : Créer une colonne 'Résultat' qui affiche 'Admis' si la moyenne >= 10, sinon 'Recalé'.",
        "Feuille Ventes : Créer un histogramme comparant les ventes totales des vendeurs.",
        "Feuille Budget : Créer un camembert montrant la répartition des dépenses réelles."
    ]
    for m in missions:
        document.add_paragraph(m, style='List Number')

    # --- Quiz & Evaluation ---
    document.add_heading('4. Grille d\'Évaluation (Note / 20)', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Critère'
    hdr_cells[1].text = 'Points'

    data = [
        ("Fonction SI correcte (Pas d'erreur de syntaxe)", "5"),
        ("Choix du graphique pertinent (Barres pour Ventes, Camembert pour Budget)", "5"),
        ("Présence de Titres sur les graphiques", "5"),
        ("Lisibilité (Légendes, Étiquettes de données)", "5")
    ]

    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]

    filename = "Brief_Pedagogique_S4.docx"
    document.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    create_brief_S4()
