from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_brief_S5():
    document = Document()

    # --- Title ---
    title = document.add_heading('Brief Pédagogique Semaine 5 : Intro BI', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Résumé ---
    document.add_heading('1. Résumé de la Semaine', level=1)
    p = document.add_paragraph()
    p.add_run('Transition critique : l\'étudiant comprend les limites d\'Excel (lenteur, statique) et découvre la puissance de l\'interactivité via un outil de Business Intelligence (Power BI ou Looker Studio).')

    # --- Compétences Visées ---
    document.add_heading('2. Compétences Visées', level=1)
    skills = [
        "Comprendre la différence entre un Tableur (Saisie/Calcul) et un Outil BI (Visualisation/Décision).",
        "Connecter une source de données (Excel) à un outil BI.",
        "Créer des visuels interactifs (Filtrage croisé).",
        "Utiliser des visuels avancés (Cartes géographiques)."
    ]
    for skill in skills:
        document.add_paragraph(skill, style='List Bullet')

    # --- Projet Pratique ---
    document.add_heading('3. Description du Projet Pratique', level=1)
    document.add_paragraph('Scénario : "Le Grand Rapport 2023"')
    document.add_paragraph('L\'étudiant manipule un fichier plus volumineux "dataset_S5_bi.xlsx" (1000 lignes).')

    document.add_heading('Mission :', level=2)
    missions = [
        "Importer le fichier Excel dans Power BI Desktop (Windows) ou Looker Studio (Mac/Web).",
        "Créer un Histogramme (Ventes par Vendeur) et un Anneau (Ventes par Catégorie).",
        "Tester l'interactivité : Cliquer sur 'Informatique' doit filtrer les vendeurs.",
        "Ajouter une Carte géographique des ventes par Région."
    ]
    for m in missions:
        document.add_paragraph(m, style='List Number')

    # --- Quiz & Evaluation ---
    document.add_heading('4. Grille d\'Évaluation (Note / 20)', level=1)

    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Critère'
    hdr_cells[1].text = 'Points'

    data = [
        ("Connexion réussie à la source de données", "5"),
        ("Création de 3 visuels différents (Barre, Anneau, Carte)", "10"),
        ("Démonstration de l'interactivité (Filtrage croisé)", "5")
    ]

    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]

    filename = "Brief_Pedagogique_S5.docx"
    document.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    create_brief_S5()
