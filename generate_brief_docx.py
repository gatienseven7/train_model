from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_brief_docx():
    document = Document()

    # --- Title ---
    title = document.add_heading('Brief Pédagogique : Atelier "Data Expert & SQL"', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph('Durée : 3h30 | Niveau : Intermédiaire à Expert')

    # --- Résumé ---
    document.add_heading('1. Résumé de l\'Atelier', level=1)
    p = document.add_paragraph()
    p.add_run('Cet atelier intensif transforme les participants en experts de la manipulation de données sur Excel (Tri/Filtre avancé, Power Query, VBA) et Google Sheets (QUERY), avant de plonger dans les bases de données SQL (Jointures) et NoSQL (Documents).')

    # --- Compétences Visées ---
    document.add_heading('2. Compétences Visées', level=1)
    skills = [
        "Maîtriser le Tri Multi-critères et les Filtres Avancés (Criteria Range).",
        "Automatiser le nettoyage de données avec Power Query (Langage M) et VBA.",
        "Croiser des données issues de multiples sources (VLOOKUP, Merge Queries).",
        "Exécuter des requêtes SQL complexes (JOIN) et NoSQL (Imbrication).",
        "Comprendre la différence structurelle entre Table Relationnelle et Document JSON."
    ]
    for skill in skills:
        document.add_paragraph(skill, style='List Bullet')

    # --- Projet Pratique ---
    document.add_heading('3. Description du Projet Pratique', level=1)
    document.add_paragraph('Scénario : "Telco-Data-Rescue"')
    document.add_paragraph('Les étudiants reçoivent un fichier "dataset_clients_advanced.xlsx" contenant 3 feuilles (Clients, Offres, Logs). Les données sont "sales" (dates mixtes, doublons, incohérences).')

    document.add_heading('Mission :', level=2)
    missions = [
        "Nettoyer les données (supprimer doublons, normaliser dates).",
        "Enrichir la table Clients avec les Prix des offres (via VLOOKUP ou Merge).",
        "Créer un formulaire de saisie automatisé via VBA.",
        "Exporter les données propres vers une base SQL en ligne et effectuer des requêtes de statistiques."
    ]
    for m in missions:
        document.add_paragraph(m, style='List Number')

    # --- Grille d'Évaluation ---
    document.add_heading('4. Grille d\'Évaluation (Note / 100)', level=1)

    table = document.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Critère'
    hdr_cells[1].text = 'Indicateur de Réussite (KPI)'
    hdr_cells[2].text = 'Points'

    data = [
        ("Nettoyage Avancé", "Dates normalisées (AAAA-MM-JJ) via Power Query M.", "20"),
        ("Croisement de Données", "La colonne 'Prix' est correctement ajoutée (VLOOKUP ou Merge).", "20"),
        ("Filtres Complexes", "Extraction correcte des 'Clients IDF en Retard' sur une autre feuille.", "15"),
        ("Automatisation (VBA)", "Le bouton 'Ajouter Client' fonctionne et ajoute une ligne.", "15"),
        ("SQL (Jointures)", "Requête JOIN fonctionnelle entre Clients et Offres.", "15"),
        ("Google Sheets", "Utilisation correcte de la fonction =QUERY().", "15")
    ]

    for item in data:
        row_cells = table.add_row().cells
        row_cells[0].text = item[0]
        row_cells[1].text = item[1]
        row_cells[2].text = item[2]

    # --- Guide de Correction ---
    document.add_heading('5. Guide de Correction Détaillé (Pour Instructeurs)', level=1)
    document.add_paragraph('Utilisez ce guide pour valider les rendus des étudiants.')

    document.add_heading('A. Vérification Excel (Fichier .xlsx)', level=2)
    checks_excel = [
        "Ouvrir Power Query (Données > Requêtes) : Vérifier qu'il y a au moins une étape 'Changed Type' avec Locale=US/FR pour les dates.",
        "Vérifier la feuille 'Clients_Nettoyés' : Aucune ligne vide, aucun doublon d'ID.",
        "Test VBA : Cliquer sur le bouton. Une ligne doit s'ajouter en bas du tableau.",
        "Formules : Vérifier la présence de VLOOKUP ou XLOOKUP dans la colonne Prix."
    ]
    for c in checks_excel:
        document.add_paragraph(c, style='List Bullet')

    document.add_heading('B. Vérification SQL (Capture d\'écran ou Lien Fiddle)', level=2)
    checks_sql = [
        "Le script doit contenir un CREATE TABLE pour Clients et Offres.",
        "La requête SELECT doit avoir un 'JOIN Offres ON...'.",
        "Le résultat ne doit pas montrer de NULL dans la colonne Prix."
    ]
    for c in checks_sql:
        document.add_paragraph(c, style='List Bullet')

    filename = "Brief_Pedagogique.docx"
    document.save(filename)
    print(f"Document saved as {filename}")

if __name__ == "__main__":
    create_brief_docx()
