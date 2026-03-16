import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading

def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def main():
    doc = Document()

    # Title
    title = doc.add_heading('PROJET FIL ROUGE : KADEA TELCO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, '"L\'enfer des requêtes et de la rétention"\n', bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, 'Historique et Narration Globale', level=1)
    doc.add_paragraph(
        "Le Pitch Global : Vous êtes Data Analyst chez Kadea Telco. Face à une concurrence féroce et "
        "une instabilité réseau sans précédent, la survie de l'entreprise repose sur votre capacité à manipuler, "
        "croiser et soumettre des milliers de données brutes à des règles de gestion d'une extrême complexité."
    )

    # Phase 1
    add_heading(doc, 'PHASE 1 : Le chaos des Logs et le défi du croisement', level=2)
    doc.add_paragraph(
        "Vous arrivez dans un contexte tendu. Les ingénieurs réseau ont extrait des centaines de milliers de lignes "
        "de 'logs' bruts contenant uniquement des identifiants d'antennes (Cell_ID) et des heures de pannes. Il manque "
        "l'essentiel : la localisation et le type d'équipement. Le directeur technique vous remet un vieux registre "
        "de maintenance mal entretenu. Pour que vos données aient un sens, vous devez réconcilier ces deux mondes."
    )
    doc.add_paragraph("- Objectif : Relier des tables hétérogènes sans erreur de correspondance.")
    doc.add_paragraph("- Outils : RECHERCHEV / VLOOKUP (Excel) / QUERY (Google Sheets).")
    doc.add_paragraph("- Livrable : Une base unifiée traduisant les codes obscurs en adresses réelles.")
    doc.add_paragraph("- Compétences : Manipulation de matrices, gestion des erreurs (#N/A).")

    # Phase 2
    add_heading(doc, "PHASE 2 : L'offensive et la synthèse macroscopique", level=2)
    doc.add_paragraph(
        "L'offensive du concurrent est brutale. Le management est submergé par la granularité de vos millions de lignes ; "
        "ils ne peuvent pas lire l'information au cas par cas. Il leur faut une vision 'hélicoptère' pour décider où "
        "envoyer les équipes de réparation en priorité. Vous devez transformer cette montagne de données en statistiques globales infaillibles."
    )
    doc.add_paragraph("- Objectif : Passer de la donnée unitaire (micro) à la statistique agrégée (macro).")
    doc.add_paragraph("- Outils : Tableaux Croisés Dynamiques (TCD), GROUP BY, SUM, AVG.")
    doc.add_paragraph("- Livrable : Un tableau de synthèse multicritères par Région et Technologie.")
    doc.add_paragraph("- Compétences : Fonctions d'agrégation, analyse de zones critiques.")

    # Phase 3
    add_heading(doc, 'PHASE 3 : Le labyrinthe de la Rétention', level=2)
    doc.add_paragraph(
        "Le Directeur Général est furieux : les clients partent à la concurrence (Churn). Pour éteindre l'incendie, "
        "un 'Plan de Rétention' d'une complexité inouïe est décidé. C'est à vous de le coder. Vous allez devoir traduire "
        "des règles humaines nuancées en une matrice logique implacable. Entre l'ancienneté, la durée de panne et le type de forfait, "
        "chaque erreur de formule pourrait coûter des millions à l'entreprise."
    )
    doc.add_paragraph("- Objectif : Traduire une règle métier complexe en formule mathématique infaillible.")
    doc.add_paragraph("- Outils : Fonctions SI (IF) imbriquées, croisées avec des opérateurs ET (AND) et OU (OR).")
    doc.add_paragraph("- Livrable : Une colonne 'Taux de Remise' 100% automatisée.")
    doc.add_paragraph("- Compétences : Logique algorithmique avancée, calcul d'impact financier.")

    # Phase 4
    add_heading(doc, 'PHASE 4 : La libération via la Business Intelligence', level=2)
    doc.add_paragraph(
        "Vous avez survécu à l'enfer des formules, mais votre fichier Excel est devenu une 'usine à gaz' trop lourde. "
        "Le CODIR exige désormais un pilotage visuel, fluide et en temps réel. C'est le moment d'abandonner le tableur pour l'industrialisation. "
        "Vous devez migrer toute votre intelligence métier vers un outil de Business Intelligence pour transformer vos calculs en un cockpit interactif."
    )
    doc.add_paragraph("- Objectif : Industrialiser la donnée : passer du fichier statique au Dashboard dynamique.")
    doc.add_paragraph("- Outils : Migration vers Power BI.")
    doc.add_paragraph("- Livrable : Un tableau de bord complet avec segments dynamiques et cartes géographiques.")
    doc.add_paragraph("- Compétences : Processus ETL (Importation), modélisation de données, Design de visuels interactifs.")

    doc.add_page_break()

    # Grading Rubric
    add_heading(doc, 'Grille de Notation (100 Points)', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Phase'
    hdr_cells[1].text = 'Critères'
    hdr_cells[2].text = 'Points (Max 25/Phase)'

    row = table.add_row().cells
    row[0].text = 'Phase 1 : Nettoyage et Croisement'
    row[1].text = 'Gestion des doublons, valeurs manquantes, RECHERCHEV correct, gestion des erreurs #N/A'
    row[2].text = '25'

    row = table.add_row().cells
    row[0].text = 'Phase 2 : Synthèse Macroscopique'
    row[1].text = 'Création correcte des TCD, agrégations exactes (Somme, Moyenne), filtres pertinents'
    row[2].text = '25'

    row = table.add_row().cells
    row[0].text = 'Phase 3 : Logique de Rétention'
    row[1].text = 'Fonctions SI imbriquées sans erreur, utilisation de ET/OU, calcul du taux de remise exact'
    row[2].text = '25'

    row = table.add_row().cells
    row[0].text = 'Phase 4 : Business Intelligence (PowerBI)'
    row[1].text = 'Importation réussie (ETL), relations de tables correctes, visuels interactifs et clairs'
    row[2].text = '25'

    doc.add_page_break()

    # Correction Guides
    add_heading(doc, 'Guide de Correction et Méthodologie', level=1)

    add_heading(doc, 'Méthode 1 : Microsoft Excel', level=2)
    doc.add_paragraph("1. Power Query : Idéal pour nettoyer les données (Supprimer les doublons, remplacer les valeurs nulles).")
    doc.add_paragraph("2. RECHERCHEV : =RECHERCHEV(A2; 'Registre_Maintenance'!A:F; 2; FAUX)")
    doc.add_paragraph("3. VBA (Optionnel) : Utiliser une macro pour automatiser la suppression des lignes en erreur et générer les TCD.")
    doc.add_paragraph("4. SI Imbriqués : =SI(ET(Anciennete>12; Forfait=\"Premium\"); 0,2; SI(OU(Plaintes>3; Duree>120); 0,1; 0))")

    add_heading(doc, 'Méthode 2 : Google Sheets', level=2)
    doc.add_paragraph("1. QUERY / Data Clean-up : Utiliser Data > Data cleanup > Remove duplicates.")
    doc.add_paragraph("2. QUERY / VLOOKUP : =VLOOKUP(A2, Registre_Maintenance!A:F, 2, FALSE)")
    doc.add_paragraph("3. AppScript (Optionnel) : Script pour fusionner les feuilles ou automatiser la création de rapports PDF.")
    doc.add_paragraph("4. ARRAYFORMULA avec IF : =ARRAYFORMULA(IF(AND(B2:B>12, C2:C=\"Premium\"), 0.2, IF(OR(F2:F>3, D2:D>120), 0.1, 0))) *Note: AND/OR require special handling in ArrayFormulas in GSheets, encourage standard drag-down IFs or nested IFs.*")

    doc.save('Guide_Instructeur.docx')
    print("Guide_Instructeur.docx generated.")

if __name__ == '__main__':
    main()
