from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Couleurs
KADEA_RED = RGBColor(237, 28, 36)
ANTHRACITE = RGBColor(40, 40, 40)

def add_heading(doc, text, level=1, color=KADEA_RED):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return heading

def add_paragraph(doc, text, bold=False, color=ANTHRACITE):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    return p

def create_instructor_brief():
    doc = Document()

    # Titre
    title = add_heading(doc, 'Brief Instructeur - Module 2 : Modélisation Avancée et Solutions Décisionnelles', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_paragraph(doc, 'WEEK 1 | Relational Database Foundations & Schema Design', bold=True)
    add_paragraph(doc, 'Durée : 120 Minutes | Ton : Smart-Fun (Geek) | Code Couleur : Minimalisme Premium')

    doc.add_page_break()

    # Introduction
    add_heading(doc, 'Objectif de la Session')
    add_paragraph(doc, "Transiter de la logique limitante de 'Tableur' vers la puissance des systèmes de gestion de bases de données relationnelles (RDBMS).")
    add_paragraph(doc, "Key Milestone : Capacité à lire, comprendre et concevoir un Database Schema relationnel normalisé.")

    # Bloc 1
    add_heading(doc, '🔴 BLOC 1 : CORE CONCEPTS (30 min) - Le choc des Titans')
    add_paragraph(doc, 'Objectif :', bold=True)
    add_paragraph(doc, 'Faire comprendre POURQUOI on abandonne les vieilles habitudes sur Excel.')
    add_paragraph(doc, 'Le Discours (Storytelling) :', bold=True)
    add_paragraph(doc, "« Jusqu'ici, vous étiez des artisans sur Excel. Mais le tableur a un plafond de verre : le Excel Ceiling. À partir d'un million de lignes, le logiciel suffoque. Bienvenue dans l'ingénierie de la donnée. Aujourd'hui, nous passons aux RDBMS (Relational Database Management Systems). »")
    add_paragraph(doc, 'Le Discours (Data Integrity & Scalability) :', bold=True)
    add_paragraph(doc, "« Pourquoi les banques n'utilisent pas Excel ? Pour la Data Integrity et la Scalability. Un RDBMS applique des règles strictes (ex: interdiction des doublons, atomicité des transactions) empêchant toute corruption. »")

    # Bloc 2
    add_heading(doc, '🔴 BLOC 2 : TECHNICAL VOCAB (30 min) - Parler le langage machine')
    add_paragraph(doc, 'Le Discours (Tables, Fields & Records) :', bold=True)
    add_paragraph(doc, "« Oubliez le jargon bureautique. Dans le modèle relationnel, l'information est dans des Tables. Chaque ligne est un Record. Les colonnes sont des Fields. »")
    add_paragraph(doc, 'Le Discours (PK & FK) :', bold=True)
    add_paragraph(doc, "« La Primary Key (PK) est l'identifiant absolu et unique de chaque Record. La Foreign Key (FK) pointe vers la PK de la table parente pour créer le lien. C'est le ciment de toute base de données ! »")

    # Bloc 3
    add_heading(doc, '🔴 BLOC 3 : WORKSHOP (30 min) - Les mains dans le cambouis')
    add_paragraph(doc, 'Setup de l\'environnement (Step-by-Step) :', bold=True)
    add_paragraph(doc, "1. Demandez aux étudiants d'installer DBeaver ou SQLite Browser.\n2. [Placeholder Image : Capture interface anglaise DBeaver] (Expliquer l'interface en français).\n3. Créez une base vierge 'Kadea_Lab.db'.")
    add_paragraph(doc, 'Initialisation et Crash Test :', bold=True)
    add_paragraph(doc, "- Créer table 'Customers' (PK: ID client).\n- Créer table 'Orders' (FK: ID client).\n- Test d'intégrité : Faites-leur insérer une commande pour un client inexistant pour déclencher une erreur et prouver la Data Integrity.")

    # Bloc 4 (Kadea Telco Context)
    add_heading(doc, '🔴 BLOC 4 : EXERCISE (30 min) - Design Thinking (Contexte Kadea Telco)')
    add_paragraph(doc, 'Le Discours (CDM/MCD) :', bold=True)
    add_paragraph(doc, "« Un Data Analyst modélise toujours sur papier via un Conceptual Data Model (CDM). Abstrayons le monde réel ! »")
    add_paragraph(doc, 'La Mission (Contexte Adapté) :', bold=True)
    add_paragraph(doc, "Les étudiants doivent modéliser l'infrastructure réseau de Kadea Telco sur les 26 provinces de la RDC. Entités : Provinces, Cell_Towers (Antennes), Technicians.")
    add_paragraph(doc, 'Le Piège Formateur (Many-to-Many) :', bold=True)
    add_paragraph(doc, "Un technicien répare plusieurs antennes, et une antenne est maintenue par plusieurs techniciens. Solution attendue : création d'une table intermédiaire (ex: 'Maintenance_Logs' ou 'Assignments').")

    doc.add_page_break()

    # NOUVEAU BLOC : ADVANCED ARCHITECTURE & OPTIMIZATION (High-Level / Senior Trainer)
    add_heading(doc, '🔴 BLOC 5 : ADVANCED ARCHITECTURE (Masterclass Formateur)')
    add_paragraph(doc, 'Objectif :', bold=True)
    add_paragraph(doc, 'Transmettre une vision "Senior" en allant au-delà de la simple syntaxe SQL. Aborder la mécanique interne du SGBDR pour concevoir des systèmes performants à grande échelle.')

    add_paragraph(doc, 'Le Discours (Normalisation vs Dénormalisation) :', bold=True)
    add_paragraph(doc, "« Une base de données opérationnelle (OLTP) se doit d'être normalisée, idéalement en 3ème Forme Normale (3NF), pour éviter toute redondance. Mais attention, en Business Intelligence (OLAP), nous ferons délibérément l'inverse : la dénormalisation, pour réduire le coût des jointures lors de la lecture de millions de lignes. »")

    add_paragraph(doc, 'Le Discours (Performance et Indexation) :', bold=True)
    add_paragraph(doc, "« Si votre requête est lente, le problème vient souvent de l\'optimiseur de requêtes (Query Planner) qui effectue un Full Table Scan. Comment l'éviter ? En créant des Index (B-Tree pour les plages de dates, Hash pour les correspondances exactes). Mais attention, un index accélère la lecture (SELECT) au détriment de l'écriture (INSERT/UPDATE). C'est un arbitrage d'architecte. »")

    add_paragraph(doc, 'Le Discours (SQL Avancé - Window Functions) :', bold=True)
    add_paragraph(doc, "« Le SQL n\'est pas mort, il a évolué. Pour des analyses statistiques poussées sans perdre le grain de vos données, oubliez le simple GROUP BY. Les CTEs (Common Table Expressions avec la clause WITH) rendent votre code lisible, et les Window Functions (OVER, PARTITION BY, LAG, LEAD) vous permettent de calculer des moyennes mobiles temporelles avec une élégance redoutable. »")

    doc.add_page_break()

    # Grille d'évaluation
    add_heading(doc, 'Grille d\'évaluation (Grading Rubric) - Exercice Bloc 4')

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Critère'
    hdr_cells[1].text = 'Points'
    hdr_cells[2].text = 'Description pour obtenir les points'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Identification des Entités'
    row_cells[1].text = '3 pts'
    row_cells[2].text = 'Les 3 entités (Provinces, Cell_Towers, Technicians) sont clairement identifiées comme des Tables.'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Attribution des Primary Keys (PK)'
    row_cells[1].text = '3 pts'
    row_cells[2].text = 'Chaque table possède un Field dédié comme identifiant unique (ex: Province_ID).'

    row_cells = table.add_row().cells
    row_cells[0].text = 'Gestion de la relation One-to-Many'
    row_cells[1].text = '2 pts'
    row_cells[2].text = "La table 'Cell_Towers' contient une Foreign Key (FK) pointant vers 'Province_ID'."

    row_cells = table.add_row().cells
    row_cells[0].text = 'Résolution du piège Many-to-Many'
    row_cells[1].text = '2 pts'
    row_cells[2].text = 'Création correcte d\'une table intermédiaire liant Cell_Towers et Technicians via deux FKs.'

    add_paragraph(doc, '\nTotal : 10 points', bold=True)

    doc.save('Module2_Week1_Instructor_Brief.docx')
    print("Guide instructeur 'Module2_Week1_Instructor_Brief.docx' généré avec succès.")

if __name__ == "__main__":
    create_instructor_brief()
